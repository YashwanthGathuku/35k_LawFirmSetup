"""
Tegifa Legal — The Autonomous Shadow Redliner
Uses python-docx and LLMs to locate conflicting clauses in Word documents
and automatically generate a new redlined document with harmonized language.
"""
import os
import logging
from docx import Document
from docx.shared import RGBColor, Pt

logger = logging.getLogger("tegifa.redliner")


def find_and_redline_document(
    file_path: str,
    target_conflict_term: str,
    harmonization_instructions: str,
    llm
) -> str:
    """
    Reads a .docx file, locates the paragraph containing the conflict,
    asks the LLM to rewrite it based on instructions, and generates
    a redlined output file.
    
    Returns the path to the redlined .docx file.
    """
    if not file_path.endswith(".docx"):
        raise ValueError("Shadow Redliner currently only supports .docx files.")
        
    doc = Document(file_path)
    
    # 1. Locate the paragraph
    target_paragraph = None
    original_text = ""
    for para in doc.paragraphs:
        if target_conflict_term.lower() in para.text.lower():
            target_paragraph = para
            original_text = para.text
            break
            
    if not target_paragraph:
        raise ValueError(f"Could not find the term '{target_conflict_term}' in the document.")

    # 2. Ask LLM to rewrite
    logger.info("Shadow Redliner: Drafting harmonization for '%s'", target_conflict_term)
    prompt = (
        "You are an expert Legal Drafter.\n"
        "Rewrite the following contract clause to resolve a conflict.\n"
        f"Original Clause: {original_text}\n"
        f"Instructions: {harmonization_instructions}\n\n"
        "Return ONLY the exact rewritten paragraph text. Do not include introductory text."
    )
    new_text = str(llm.complete(prompt)).strip()
    
    # 3. Apply Visual Redline
    # We clear the existing paragraph and add our redline simulation
    target_paragraph.clear()
    
    # Deleted text (Red Strikethrough)
    run_del = target_paragraph.add_run(original_text)
    run_del.font.strike = True
    run_del.font.color.rgb = RGBColor(255, 0, 0)
    
    # Add a space
    target_paragraph.add_run(" ")
    
    # Inserted text (Green Bold)
    run_ins = target_paragraph.add_run(new_text)
    run_ins.font.bold = True
    run_ins.font.color.rgb = RGBColor(0, 176, 80)
    
    # 4. Save the new document
    base_name = os.path.basename(file_path)
    dir_name = os.path.dirname(file_path)
    redline_name = f"Redlined_{base_name}"
    out_path = os.path.join(dir_name, redline_name)
    
    doc.save(out_path)
    logger.info("Saved redlined document to %s", out_path)
    
    return out_path
