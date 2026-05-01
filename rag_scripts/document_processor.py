"""
Tegifa Legal — Advanced Document Processor
Handles intelligent document ingestion with:
- Nougat OCR for scanned/complex PDFs (primary extraction method)
- Page-level metadata tracking through the entire pipeline
- Table detection and structured extraction
- Legal clause tagging at the chunk level
"""
import os
import re
import logging
from dataclasses import dataclass, field
from typing import Optional

logger = logging.getLogger("tegifa.document_processor")


# ---------------------------------------------------------------------------
# Legal Clause Patterns (used during ingestion to pre-tag chunks)
# ---------------------------------------------------------------------------
CLAUSE_PATTERNS = {
    "termination": [
        r"terminat(?:e|ion|ing)",
        r"cancel(?:lation)?",
        r"expir(?:e|ation|y)",
        r"right to end",
    ],
    "indemnification": [
        r"indemnif(?:y|ication|ied)",
        r"hold\s+harmless",
        r"defend and indemnify",
    ],
    "liability": [
        r"liabilit(?:y|ies)",
        r"limitation of liability",
        r"consequential damages",
        r"aggregate liability",
        r"shall not be liable",
        r"cap on damages",
    ],
    "confidentiality": [
        r"confidential(?:ity)?",
        r"non-disclosure",
        r"proprietary information",
        r"trade secret",
        r"NDA",
    ],
    "intellectual_property": [
        r"intellectual property",
        r"\bIP\b",
        r"patent",
        r"copyright",
        r"trademark",
        r"work.?for.?hire",
        r"assignment of (?:rights|inventions)",
    ],
    "governing_law": [
        r"governing law",
        r"choice of law",
        r"jurisdiction",
        r"venue",
        r"dispute resolution",
        r"arbitrat(?:e|ion)",
    ],
    "payment": [
        r"payment\s+terms?",
        r"invoice",
        r"compensation",
        r"fee(?:s)?",
        r"remuneration",
        r"late\s+(?:payment|fee)",
        r"net\s+\d+",
    ],
    "force_majeure": [
        r"force majeure",
        r"act of god",
        r"unforeseeable",
        r"beyond.{0,20}control",
    ],
    "non_compete": [
        r"non-?compet(?:e|ition)",
        r"restrictive covenant",
        r"solicitation",
        r"non-?solicit",
    ],
    "warranty": [
        r"warrant(?:y|ies|s)",
        r"representation(?:s)?",
        r"as[- ]is",
        r"merchantability",
        r"fitness for.{0,20}purpose",
    ],
    "insurance": [
        r"insurance",
        r"coverage",
        r"policy limit",
        r"certificate of insurance",
    ],
    "assignment": [
        r"assign(?:ment|ability)",
        r"transfer(?:ability)?",
        r"delegation",
        r"successor",
    ],
    "severability": [
        r"severab(?:le|ility)",
        r"invalid(?:ity)?.*provision",
        r"unenforceable",
    ],
    "entire_agreement": [
        r"entire agreement",
        r"merger clause",
        r"integration clause",
        r"supersede(?:s)?",
    ],
}


@dataclass
class ProcessedChunk:
    """A single chunk of text with full metadata for citation and analysis."""

    text: str
    file_name: str
    file_path: str
    page_number: Optional[int] = None
    paragraph_index: Optional[int] = None
    total_pages: Optional[int] = None
    chunk_index: int = 0
    total_chunks: int = 0
    detected_clauses: list = field(default_factory=list)
    contains_table: bool = False
    section_heading: Optional[str] = None
    char_offset_start: int = 0
    char_offset_end: int = 0


@dataclass
class ProcessedDocument:
    """A fully processed document with all chunks and metadata."""

    file_name: str
    file_path: str
    total_pages: int
    total_chunks: int
    chunks: list  # list[ProcessedChunk]
    detected_tables: int = 0
    document_type: str = "unknown"  # contract, brief, statute, memo, etc.


def detect_clauses(text: str) -> list[str]:
    """Detect which legal clause types are present in a text chunk."""
    text_lower = text.lower()
    found = []
    for clause_type, patterns in CLAUSE_PATTERNS.items():
        for pattern in patterns:
            if re.search(pattern, text_lower):
                found.append(clause_type)
                break  # One match per clause type is enough
    return found


def detect_table(text: str) -> bool:
    """Heuristic table detection in text content."""
    indicators = [
        # Markdown-style table separators
        re.search(r"\|[\s\-]+\|", text),
        # Tab-separated columns (3+ tabs on a line)
        re.search(r"\t.*\t.*\t", text),
        # Repeated aligned whitespace patterns
        len(re.findall(r"^\s{2,}\S+\s{2,}\S+", text, re.MULTILINE)) >= 3,
    ]
    return any(indicators)


def detect_section_heading(text: str) -> Optional[str]:
    """Extract section heading if the chunk starts with one."""
    patterns = [
        # "Section 1.2 — Title" or "ARTICLE III"
        r"^(?:Section|Article|Clause|Part)\s+[\dIVXivx]+[\.\):]?\s*[—\-:]?\s*(.+)",
        # Numbered heading: "1.2 Title" or "3. Title"
        r"^(\d+[\.\d]*)\s+([A-Z][A-Za-z\s]{3,50})$",
        # ALL CAPS heading
        r"^([A-Z][A-Z\s]{5,60})$",
    ]
    first_line = text.strip().split("\n")[0].strip()
    for pattern in patterns:
        match = re.match(pattern, first_line)
        if match:
            return first_line[:80]
    return None


def classify_document_type(full_text: str) -> str:
    """Classify what type of legal document this is."""
    text_lower = full_text[:5000].lower()  # Check first 5K chars

    type_signals = {
        "contract": [
            r"agreement",
            r"parties hereto",
            r"whereas",
            r"now,?\s+therefore",
            r"in witness whereof",
            r"executed as of",
        ],
        "brief": [
            r"court of",
            r"plaintiff",
            r"defendant",
            r"motion for",
            r"memorandum of law",
            r"respectfully submitted",
        ],
        "statute": [
            r"be it enacted",
            r"§\s*\d+",
            r"public law",
            r"u\.?s\.?c\.?",
            r"chapter \d+",
        ],
        "memorandum": [
            r"^memo(?:randum)?",
            r"to:\s+",
            r"from:\s+",
            r"re:\s+",
            r"subject:\s+",
        ],
        "complaint": [
            r"comes now",
            r"causes? of action",
            r"prayer for relief",
            r"jury (?:trial )?demand",
        ],
    }

    scores = {}
    for doc_type, patterns in type_signals.items():
        score = sum(1 for p in patterns if re.search(p, text_lower))
        if score > 0:
            scores[doc_type] = score

    if scores:
        return max(scores, key=scores.get)
    return "unknown"


def extract_pages_from_pdf(filepath: str, use_nougat: bool = True) -> list[dict]:
    """
    Extract text from PDF with page-level tracking.

    Strategy:
    1. Try Nougat OCR first (best for scanned docs, tables, formulas)
    2. Fall back to pypdf page-by-page extraction (preserves page numbers)
    3. Fall back to LlamaIndex SimpleDirectoryReader

    Returns list of {"text": ..., "page_number": ..., "total_pages": ...}
    """
    pages = []

    # Strategy 1: Nougat OCR
    if use_nougat:
        try:
            from nougat import NougatModel
            from nougat.utils.checkpoint import get_checkpoint
            import torch

            logger.info("Attempting Nougat OCR for: %s", os.path.basename(filepath))
            checkpoint = get_checkpoint("nougat-base")
            model = NougatModel.from_pretrained(checkpoint)
            model.eval()

            if torch.cuda.is_available():
                model = model.to("cuda")

            from nougat.utils.dataset import LazyDataset
            from torch.utils.data import DataLoader

            dataset = LazyDataset(filepath, model.encoder.prepare_input)
            dataloader = DataLoader(dataset, batch_size=1, shuffle=False)

            total = len(dataset)
            for page_idx, sample in enumerate(dataloader):
                if torch.cuda.is_available():
                    sample = sample.to("cuda")
                output = model.inference(image_tensors=sample)
                page_text = output["predictions"][0] if output["predictions"] else ""
                if page_text and page_text != "[MISSING_PAGE_EMPTY]":
                    pages.append({
                        "text": page_text,
                        "page_number": page_idx + 1,
                        "total_pages": total,
                    })

            if pages:
                logger.info("Nougat extracted %d pages from %s", len(pages), os.path.basename(filepath))
                return pages
            else:
                logger.warning("Nougat returned empty. Falling back to pypdf.")

        except ImportError:
            logger.info("Nougat not installed. Falling back to pypdf.")
        except Exception as e:
            logger.warning("Nougat failed: %s. Falling back to pypdf.", e)

    # Strategy 2: pypdf page-by-page (preserves page numbers)
    try:
        from pypdf import PdfReader

        reader = PdfReader(filepath)
        total = len(reader.pages)

        for page_idx, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                pages.append({
                    "text": text,
                    "page_number": page_idx + 1,
                    "total_pages": total,
                })

        if pages:
            logger.info("pypdf extracted %d pages from %s", len(pages), os.path.basename(filepath))
            return pages

    except ImportError:
        logger.warning("pypdf not installed.")
    except Exception as e:
        logger.warning("pypdf failed: %s", e)

    # Strategy 3: LlamaIndex SimpleDirectoryReader (no page tracking)
    try:
        from llama_index.core import SimpleDirectoryReader

        docs = SimpleDirectoryReader(input_files=[filepath]).load_data()
        for doc in docs:
            pages.append({
                "text": doc.text,
                "page_number": doc.metadata.get("page_label"),
                "total_pages": len(docs),
            })
        logger.info("SimpleDirectoryReader extracted %d pages", len(pages))

    except Exception as e:
        logger.error("All extraction methods failed for %s: %s", filepath, e)

    return pages


def extract_text_file(filepath: str) -> list[dict]:
    """Extract text from non-PDF files."""
    try:
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
        return [{"text": text, "page_number": 1, "total_pages": 1}]
    except Exception as e:
        logger.error("Failed to read %s: %s", filepath, e)
        return []


def chunk_with_metadata(
    pages: list[dict],
    file_name: str,
    file_path: str,
    chunk_size: int = 1024,
    chunk_overlap: int = 128,
) -> ProcessedDocument:
    """
    Split pages into chunks while preserving page-level metadata.

    Each chunk retains:
    - Source page number(s)
    - Paragraph index within the page
    - Detected legal clause types
    - Table detection flag
    - Section heading if applicable
    - Character offsets for pinpoint citations
    """
    all_chunks: list[ProcessedChunk] = []
    full_text = "\n\n".join(p["text"] for p in pages)
    doc_type = classify_document_type(full_text)
    total_pages = pages[0]["total_pages"] if pages else 0
    table_count = 0
    global_char_offset = 0

    for page in pages:
        page_text = page["text"]
        page_num = page.get("page_number")

        # Split page into paragraphs (double newline or significant whitespace)
        paragraphs = re.split(r"\n\s*\n", page_text)

        for para_idx, paragraph in enumerate(paragraphs):
            paragraph = paragraph.strip()
            if not paragraph or len(paragraph) < 20:
                continue

            # If paragraph is longer than chunk_size, split further
            if len(paragraph) > chunk_size:
                # Split on sentence boundaries
                sentences = re.split(r"(?<=[.!?])\s+", paragraph)
                current_chunk = ""
                chunk_start = global_char_offset

                for sentence in sentences:
                    if len(current_chunk) + len(sentence) > chunk_size and current_chunk:
                        # Emit current chunk
                        has_table = detect_table(current_chunk)
                        if has_table:
                            table_count += 1

                        all_chunks.append(
                            ProcessedChunk(
                                text=current_chunk.strip(),
                                file_name=file_name,
                                file_path=file_path,
                                page_number=page_num,
                                paragraph_index=para_idx,
                                total_pages=total_pages,
                                detected_clauses=detect_clauses(current_chunk),
                                contains_table=has_table,
                                section_heading=detect_section_heading(current_chunk),
                                char_offset_start=chunk_start,
                                char_offset_end=chunk_start + len(current_chunk),
                            )
                        )

                        # Start new chunk with overlap
                        overlap_text = current_chunk[-chunk_overlap:] if chunk_overlap else ""
                        chunk_start = global_char_offset + len(current_chunk) - len(overlap_text)
                        current_chunk = overlap_text + " " + sentence
                    else:
                        current_chunk += (" " if current_chunk else "") + sentence

                # Emit remaining text
                if current_chunk.strip():
                    has_table = detect_table(current_chunk)
                    if has_table:
                        table_count += 1

                    all_chunks.append(
                        ProcessedChunk(
                            text=current_chunk.strip(),
                            file_name=file_name,
                            file_path=file_path,
                            page_number=page_num,
                            paragraph_index=para_idx,
                            total_pages=total_pages,
                            detected_clauses=detect_clauses(current_chunk),
                            contains_table=has_table,
                            section_heading=detect_section_heading(current_chunk),
                            char_offset_start=chunk_start,
                            char_offset_end=chunk_start + len(current_chunk),
                        )
                    )
            else:
                # Paragraph fits in one chunk
                has_table = detect_table(paragraph)
                if has_table:
                    table_count += 1

                all_chunks.append(
                    ProcessedChunk(
                        text=paragraph,
                        file_name=file_name,
                        file_path=file_path,
                        page_number=page_num,
                        paragraph_index=para_idx,
                        total_pages=total_pages,
                        detected_clauses=detect_clauses(paragraph),
                        contains_table=has_table,
                        section_heading=detect_section_heading(paragraph),
                        char_offset_start=global_char_offset,
                        char_offset_end=global_char_offset + len(paragraph),
                    )
                )

            global_char_offset += len(paragraph) + 2  # +2 for paragraph separator

    # Number chunks
    for i, chunk in enumerate(all_chunks):
        chunk.chunk_index = i
        chunk.total_chunks = len(all_chunks)

    return ProcessedDocument(
        file_name=file_name,
        file_path=file_path,
        total_pages=total_pages,
        total_chunks=len(all_chunks),
        chunks=all_chunks,
        detected_tables=table_count,
        document_type=doc_type,
    )


def process_file(
    filepath: str,
    use_nougat: bool = True,
    chunk_size: int = 1024,
    chunk_overlap: int = 128,
) -> ProcessedDocument:
    """
    Full processing pipeline for a single file.

    Returns a ProcessedDocument with all chunks tagged with:
    - Page numbers
    - Paragraph indices
    - Clause types
    - Table flags
    - Section headings
    - Character offsets
    """
    file_name = os.path.basename(filepath)
    ext = os.path.splitext(filepath)[1].lower()

    logger.info("Processing file: %s (type: %s)", file_name, ext)

    if ext == ".pdf":
        pages = extract_pages_from_pdf(filepath, use_nougat=use_nougat)
    elif ext in (".txt", ".md", ".rst"):
        pages = extract_text_file(filepath)
    elif ext in (".docx",):
        try:
            import docx

            doc = docx.Document(filepath)
            text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
            pages = [{"text": text, "page_number": 1, "total_pages": 1}]
        except ImportError:
            logger.warning("python-docx not installed. Treating as text.")
            pages = extract_text_file(filepath)
    else:
        pages = extract_text_file(filepath)

    if not pages:
        logger.warning("No content extracted from %s", file_name)
        return ProcessedDocument(
            file_name=file_name,
            file_path=filepath,
            total_pages=0,
            total_chunks=0,
            chunks=[],
        )

    doc = chunk_with_metadata(
        pages=pages,
        file_name=file_name,
        file_path=filepath,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
    )

    logger.info(
        "Processed %s: %d pages, %d chunks, %d tables, type=%s",
        file_name,
        doc.total_pages,
        doc.total_chunks,
        doc.detected_tables,
        doc.document_type,
    )

    return doc


def chunk_to_metadata(chunk: ProcessedChunk) -> dict:
    """Convert a ProcessedChunk to a flat metadata dict for vector store."""
    return {
        "file_name": chunk.file_name,
        "file_path": chunk.file_path,
        "page_number": chunk.page_number,
        "paragraph_index": chunk.paragraph_index,
        "total_pages": chunk.total_pages,
        "chunk_index": chunk.chunk_index,
        "total_chunks": chunk.total_chunks,
        "detected_clauses": ",".join(chunk.detected_clauses) if chunk.detected_clauses else "",
        "contains_table": chunk.contains_table,
        "section_heading": chunk.section_heading or "",
        "char_offset_start": chunk.char_offset_start,
        "char_offset_end": chunk.char_offset_end,
    }
